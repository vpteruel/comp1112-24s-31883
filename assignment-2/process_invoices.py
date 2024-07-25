import os
import subprocess
import sys

# Function to install required packages
def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Install required packages if not already installed

# openpyxl: A library to read/write Excel 2010 xlsx/xlsm/xltx/xltm files
try:
    import openpyxl
except ModuleNotFoundError:
    install('openpyxl')

# pandas: A powerful data analysis and manipulation library for Python
try:
    import pandas as pd
except ModuleNotFoundError:
    install('pandas')

# python-docx: A library for creating, updating, and reading Microsoft Word (.docx) files
try:
    from docx import Document
except ModuleNotFoundError:
    install('python-docx')

# tqdm: A library for adding progress bars to Python loops
try:
    from tqdm import tqdm
except ModuleNotFoundError:
    install('tqdm')

# List to store invoice data
invoices_data = []

# Directory containing the invoices
invoices_dir = 'assignment-2/invoices/'

print("Starting to process invoices...")

# Process each Word document with a progress bar
for filename in tqdm(os.listdir(invoices_dir)):
    if filename.endswith('.docx'):
        # Load the document
        doc = Document(os.path.join(invoices_dir, filename))
        
        # Extract invoice number
        invoice_id = doc.paragraphs[0].text.strip()
        
        # Extract product details and calculate total number of products
        product_count = 0
        product_lines = doc.paragraphs[1].text.strip().split('\n')
        for line in product_lines:
            if ':' in line:
                product, count = line.split(':')
                product_count += int(count)
        
        # Extract subtotal, tax, and total
        subtotal, tax, total = 0.0, 0.0, 0.0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith("SUBTOTAL:"):
                value_lines = text.strip().split('\n')
                for line in value_lines:
                    if ':' in line:
                        if line.startswith("SUBTOTAL:"):
                            subtotal = float(line.split('SUBTOTAL:')[1].split()[0])
                        elif line.startswith("TAX:"):
                            tax = float(line.split('TAX:')[1].split()[0])
                        elif line.startswith("TOTAL:"):
                            total = float(line.split('TOTAL:')[1].split()[0])
        
        # Append extracted data to the list
        invoices_data.append([invoice_id, product_count, subtotal, tax, total])

# Create a DataFrame
df = pd.DataFrame(invoices_data, columns=['Invoice ID', 'Total Products Purchased', 'Subtotal', 'Tax', 'Total'])

# Save the DataFrame to an Excel file
output_path = 'assignment-2/compiled_invoices.xlsx'
df.to_excel(output_path, index=False)

print("Invoices have been successfully processed and saved to", output_path)
